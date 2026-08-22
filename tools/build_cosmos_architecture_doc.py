from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "COSMOS_Proprietary_Computational_Engineering_Architecture.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 96, 108)
LIGHT_FILL = "F2F4F7"
MID_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, *, size: float | None = None, color: RGBColor | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    """
    Mark a table header row so long tables repeat headers across pages.
    """

    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "8", space: str = "4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.8)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.06

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.8)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.08


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("COSMOS | Proprietary Architecture Pattern | Confidential Draft")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Company Confidential - COSMOS Intellectual Property - Human review required before engineering release")
    set_run_font(run, size=8.5, color=MUTED)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("COSMOS PROPRIETARY COMPUTATIONAL ENGINEERING ARCHITECTURE")
    set_run_font(run, size=23, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Architecture pattern, intellectual-property doctrine, RAG knowledge foundation, "
        "CAD/topology generation workflow, simulation validation loop, UI/UX strategy, "
        "and engineering documentation system"
    )
    set_run_font(run, size=12.5, color=MUTED, italic=True)

    rows = [
        ("Project", "COSMOS - Cryogenic Optimization and Simulation Multiphysics Operating System"),
        ("Document Type", "Internal architecture and product-development specification"),
        ("Status", "Draft for founder/engineering review"),
        ("Prepared On", date.today().isoformat()),
        ("Ownership", "Company proprietary and confidential; not a clone of any external proprietary system"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=11, color=BLACK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    paragraph_border_bottom(rule, color="2E74B5", size="8", space="4")


def add_para(doc: Document, text: str, *, style: str | None = None,
             bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Number")


def add_spacer(doc: Document, *, after: float = 2) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(" ")
    set_run_font(run, size=1)


def add_callout(doc: Document, title: str, body: str, *, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent_dxa=120)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=INK, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=BLACK)
    add_spacer(doc)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              *, header_fill: str = LIGHT_FILL, font_size: float = 9.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths, indent_dxa=120)
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=INK, bold=True)
    for row in rows:
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=font_size, color=BLACK)
    add_spacer(doc)


def add_monospace_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent_dxa=120)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FBFCFE")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=8.6, color=INK, name="Courier New")
    add_spacer(doc)


def add_section(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def add_subsection(doc: Document, title: str) -> None:
    doc.add_heading(title, level=2)


def build_document() -> None:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "Core thesis",
        "COSMOS shall be developed as a proprietary computational engineering compiler for liquid rocket "
        "propulsion: requirements become a validated design state, evidence-bound solvers generate engineering "
        "parameters, geometry generators compile manufacturable CAD, external simulation tools challenge the "
        "result, and a human release gate controls any 3D-printable output."
    )

    add_section(doc, "1. Executive Doctrine")
    add_para(
        doc,
        "COSMOS is intended to be a desktop-first multiphysics engineering platform for rocket propulsion "
        "design, development, verification, and manufacturing preparation. Its purpose is not to imitate a "
        "single external tool, but to establish a company-owned computational engineering architecture that "
        "uses executable engineering logic, traceable knowledge, deterministic solvers, generated CAD, external "
        "simulation evidence, and formal release control."
    )
    add_para(
        doc,
        "The software shall work in a similar category to modern computational-engineering systems: the user "
        "enters mission and engine requirements, the system converts those requirements into a formal design "
        "contract, the engineering kernel calculates a candidate engine, and the geometry kernel produces CAD "
        "models that can be exported to specialist analysis and manufacturing systems. COSMOS must remain "
        "independent in source code, data schema, solver implementation, user experience, and proprietary "
        "design pattern."
    )
    add_bullets(doc, [
        "The RAG system is an evidence and traceability layer, not an unchecked design authority.",
        "The engineering solvers are deterministic, unit-safe, versioned, and independently testable.",
        "The generated CAD is parametric and evidence-bound before it becomes optimized or printable.",
        "Topology optimization proposes candidates; clean CAD reconstruction and external analysis decide release readiness.",
        "Human review and uploaded external results are mandatory before manufacturing release.",
    ])

    add_subsection(doc, "1.1 Product Identity")
    add_table(
        doc,
        ["Element", "Definition"],
        [
            ["Product name", "COSMOS - Cryogenic Optimization and Simulation Multiphysics Operating System"],
            ["Primary language", "Python for orchestration, scientific computing, data models, RAG, automation, and first-stage CAD generation"],
            ["Core category", "Computational engineering and AI-assisted propulsion design platform"],
            ["Primary output", "Traceable design state, engineering documentation, CAD/CAE packages, analysis comparison records, and manufacturing release packages"],
            ["Governance stance", "Evidence-first, solver-verified, externally reviewed, human released"],
        ],
        [2300, 7060],
    )

    add_subsection(doc, "1.2 Non-Clone Position")
    add_para(
        doc,
        "COSMOS may operate in the same broad category as other computational engineering systems, but the "
        "company shall treat COSMOS as a clean, independent implementation. No proprietary algorithms, source "
        "code, data structures, internal user flows, or confidential know-how from external systems shall be "
        "copied or reverse engineered. COSMOS shall be designed from first principles, public scientific "
        "knowledge, licensed references, company-generated validation data, and original engineering software."
    )
    add_callout(
        doc,
        "IP caution",
        "This document is an engineering and product-architecture draft, not legal advice. Formal IP assignment, "
        "employee/contractor invention agreements, copyright registration strategy, patent review, and trade-secret "
        "controls should be reviewed by qualified counsel."
    )

    add_section(doc, "2. COSMOS Proprietary Architecture Pattern")
    add_para(
        doc,
        "The proposed proprietary pattern is named TRACE-GEN: Traceable Retrieval-Augmented Computational "
        "Engineering Generation. TRACE-GEN is the COSMOS operating pattern that links every design output to "
        "a controlled requirement, verified equation, unit model, source reference, solver version, CAD feature, "
        "analysis result, and release decision."
    )
    add_table(
        doc,
        ["TRACE-GEN Stage", "Purpose", "Primary Artifacts"],
        [
            ["T - Technical Design Contract", "Convert user inputs into formal requirements, constraints, assumptions, and acceptance criteria.", "Design contract, requirement IDs, operating envelope"],
            ["R - Retrieval and Reference Binding", "Retrieve approved equations, material data, propellant data, correlations, and manufacturing rules.", "Evidence packets, citations, validity ranges"],
            ["A - Analytical Synthesis", "Run deterministic engineering solvers and derive a coherent preliminary engine state.", "Sizing results, margins, warnings"],
            ["C - CAD and Geometry Compilation", "Generate parametric and implicit geometry from engineering state.", "Feature graph, STEP/Parasolid, mesh seeds"],
            ["E - External Evaluation", "Export to independent CAE tools and import solver results for comparison.", "CFD/FEA/FEM/dynamics cases, result maps"],
            ["G - Governance Review", "Compare evidence and obtain human review before release.", "Review board, signoff record, issue log"],
            ["E - Engineering Documentation", "Generate traceable reports and design basis documents.", "Calculation package, design report, release note"],
            ["N - Next-Design Learning", "Use approved result deltas and test data to refine models without corrupting source evidence.", "Model calibration records, validation datasets"],
        ],
        [2000, 4350, 3010],
    )

    add_subsection(doc, "2.1 Unique COSMOS Differentiators")
    add_bullets(doc, [
        "Evidence-bound geometry: every CAD feature links to a requirement, solver output, equation, material rule, or manufacturing constraint.",
        "Design decision ledger: all automated and human decisions are recorded with timestamp, model version, reviewer, source evidence, and acceptance rationale.",
        "Solver disagreement cockpit: COSMOS predictions, Fluent/OpenFOAM/Nastran/Ansys/NX results, and manual reviewer comments are shown side by side.",
        "Manufacturing release passport: a design cannot export a print-release package without proof of geometry, mesh, thermal, structural, CFD, and human review gates.",
        "Equation validity guardian: equations retrieved by RAG are checked for units, dimensions, operating range, status, and source authority before solver use.",
        "Topology reconstruction discipline: optimized forms are converted back into parametric or implicit engineering geometry before final validation.",
        "Design lineage replay: a released engine can be regenerated from requirements, data versions, solver versions, and CAD generator versions.",
        "AI role separation: AI assists search, explanation, planning, comparison, and documentation; deterministic software owns calculations and geometry generation.",
    ])

    add_subsection(doc, "2.2 Architectural Non-Negotiables")
    add_bullets(doc, [
        "No undocumented equation shall enter production solvers.",
        "No RAG answer shall be used as a direct numerical authority without validation and source binding.",
        "No STL/3MF export shall be treated as release-ready unless the design passport is complete.",
        "No topology optimization output shall bypass clean geometry reconstruction.",
        "No external simulation result shall be accepted without solver version, mesh metadata, boundary conditions, convergence data, and load-case identity.",
        "No manufacturing package shall be released without human signoff.",
    ])

    add_section(doc, "3. Requirements-to-Release Workflow")
    add_monospace_block(
        doc,
        "User Requirements\n"
        "  -> Design Contract\n"
        "    -> Evidence Retrieval and Source Binding\n"
        "      -> Deterministic Sizing and Multiphysics Synthesis\n"
        "        -> Parametric CAD Generation\n"
        "          -> Topology / Shape Optimization Candidates\n"
        "            -> Clean CAD Reconstruction\n"
        "              -> External CFD / FEA / FEM / Dynamics Validation\n"
        "                -> COSMOS Comparison Cockpit\n"
        "                  -> Human Review Gate\n"
        "                    -> Manufacturing Release Passport\n"
        "                      -> STEP / CAE / 3MF-STL / Engineering Report Package"
    )
    add_para(
        doc,
        "The initial user inputs may include propellant pair, target thrust, chamber pressure, burn time, "
        "feed architecture, cooling strategy, manufacturing method, material preferences, altitude condition, "
        "safety factors, and packaging limits. COSMOS shall immediately convert these into a design contract "
        "rather than passing raw values directly into solvers."
    )
    add_table(
        doc,
        ["Input Class", "Examples", "Validation Requirement"],
        [
            ["Mission and operating point", "Thrust, burn time, altitude, duty cycle, restart count", "Finite values, units, completeness, intended regime"],
            ["Propulsion configuration", "Pressure-fed, pump-fed, gas-generator, expander, staged-combustion future options", "Supported cycle and maturity status"],
            ["Propellants", "LOX/LCH4, LOX/LH2, LOX/RP-1, future monopropellant or storable options", "Traceable property database and CEA species mapping"],
            ["Materials and manufacturing", "Inconel, copper alloys, GRCop family, additive process, machining limits", "Material source, thermal/structural limits, process capability"],
            ["Review strategy", "Internal validation, external Fluent/OpenFOAM/Nastran/Ansys review, test correlation", "Reviewer identity and required evidence package"],
        ],
        [1700, 4300, 3360],
    )

    add_section(doc, "4. Layered Software Architecture")
    add_para(
        doc,
        "The existing COSMOS layered architecture remains correct and should be extended with knowledge, "
        "CAD generation, optimization, simulation adapters, documentation, and release-governance layers. "
        "Dependencies should still point downward: UI calls backend APIs; backend orchestrates systems; "
        "systems combine physics models; physics depends only on core utilities and approved scientific "
        "libraries."
    )
    add_table(
        doc,
        ["Layer", "Responsibilities", "Forbidden Responsibilities"],
        [
            ["Core", "Units, constants, validation, logging, configuration, exceptions, numerical safety primitives", "Engineering equations, GUI, workflow policy"],
            ["Knowledge Foundation", "References, documents, equations, variables, units, constants, materials, propellants, assumptions, validity ranges", "Unverified AI-generated equations, direct solver execution"],
            ["Physics", "Thermochemistry, gas dynamics, heat transfer, fluids, materials physics, structural formulas", "Application workflows, CAD decisions, UI"],
            ["Systems", "Injector, cooling, chamber/nozzle, igniter, feed-system, structure, reliability assemblies", "GUI and persistence internals"],
            ["Backend", "Solver orchestration, optimization runs, dependency scheduling, project state, API boundaries", "Raw equation definitions without source traceability"],
            ["CAD/Geometry", "Parametric models, implicit geometry, feature graph, topology reconstruction, export packages", "Unsupported manufacturing release decisions"],
            ["Simulation Adapters", "Ansys, Fluent, OpenFOAM, Nastran, NX, Abaqus, Simcenter, Adams, Simulink, Dymola integration", "Silently modifying design intent"],
            ["GUI/UX", "Input capture, workflow navigation, results display, comparison cockpit, review boards", "Performing calculations directly"],
            ["Documentation", "Design basis, calculation packages, analysis reports, release passports, audit logs", "Untraceable or manually detached reports"],
        ],
        [1450, 4850, 3060],
        font_size=9,
    )

    add_subsection(doc, "4.1 Primary Python Technology Stack")
    add_table(
        doc,
        ["Domain", "Recommended Primary Tools", "Purpose"],
        [
            ["Core language", "Python 3.11+ initially; C++/Rust optional later", "Fast iteration, scientific ecosystem, automation, integration"],
            ["Data models", "Pydantic or dataclasses with strict validators", "Schema validation, serialization, traceability"],
            ["Numerics", "NumPy, SciPy, SymPy, Pint", "Solvers, optimization, symbolic checks, unit safety"],
            ["MDO", "OpenMDAO; Dakota/optiSLang adapters", "Multidisciplinary design optimization and uncertainty"],
            ["CAD", "CadQuery/OCP/OpenCASCADE; optional FreeCAD bridge; later C++ geometry kernels if required", "STEP/Parasolid-style parametric generation and solid modeling"],
            ["Mesh/visualization", "Gmsh, meshio, VTK/PyVista", "Mesh generation, conversion, result visualization"],
            ["RAG", "Vector database plus keyword index plus knowledge graph", "Evidence retrieval with citations and relationships"],
            ["Storage", "SQLite/DuckDB for local, PostgreSQL for team/enterprise, HDF5/Zarr for large results", "Project data, result arrays, audit trails"],
            ["Desktop UI", "PySide6/Qt with VTK/Qt 3D views; optional local web UI for advanced dashboards", "Professional engineering workbench"],
            ["Internal API", "FastAPI or typed service layer", "Decoupled UI/backend and future cloud/automation support"],
        ],
        [1850, 3550, 3960],
        font_size=9,
    )

    add_section(doc, "5. Knowledge Foundation and RAG Architecture")
    add_para(
        doc,
        "The knowledge folder is the source of the AI/RAG system. Its mission is to extract, normalize, "
        "govern, retrieve, and explain the mathematics, physics, chemistry, materials science, manufacturing "
        "knowledge, and engineering methods required to design and validate rocket engines. It must behave as "
        "a controlled technical-data system, not as a free-form chatbot memory."
    )
    add_table(
        doc,
        ["Knowledge Entity", "Required Fields", "Engineering Role"],
        [
            ["Reference", "Source ID, authors, title, publisher, year, DOI/URL/ISBN, status, license, confidence", "Establishes authority and permitted use"],
            ["Document", "Document ID, source reference, content hash, version, section map, security level, import pipeline", "Machine-readable source text"],
            ["Equation", "Expression, variables, units, dimensions, assumptions, validity range, source page/section, verification tests", "Reusable solver knowledge"],
            ["Variable", "Symbol, name, unit, dimension, bounds, role, finite/zero/negative policy, source", "Typed engineering quantity"],
            ["Unit", "Unit ID, symbol, dimension vector, conversion, offset, source, status", "Prevents unit drift"],
            ["Constant", "Value, uncertainty, CODATA/source version, exactness, units, dimensions", "Validated numerical constants"],
            ["Material", "Thermal, mechanical, fatigue, compatibility, AM parameters, source, allowables", "Design and manufacturing constraints"],
            ["Propellant", "Species mapping, composition, density, thermophysical properties, storage conditions, source", "Thermochemistry and feed-system inputs"],
            ["Subsystem", "Functional role, interfaces, constraints, related equations, test data", "Links knowledge to engine architecture"],
            ["Validation Dataset", "Experiment/test/simulation source, conditions, measured outputs, uncertainty", "Model verification and calibration"],
        ],
        [1550, 4500, 3310],
        font_size=8.8,
    )

    add_subsection(doc, "5.1 RAG Retrieval Rules")
    add_bullets(doc, [
        "RAG responses must cite approved references and document sections.",
        "Equation retrieval must include validity range, variables, units, dimensional signature, assumptions, and verification status.",
        "Conflicting sources must be surfaced as conflicts, not averaged silently.",
        "The AI may propose candidate design logic, but production solvers must use approved equation objects and test coverage.",
        "Every retrieved artifact must be access-controlled according to its license and company security level.",
        "Generated summaries must never overwrite source truth; they are derived artifacts linked to source hashes.",
    ])

    add_subsection(doc, "5.2 Ingestion and Approval Pipeline")
    add_numbers(doc, [
        "Acquire source: textbook, NASA report, NIST data, standard, internal test report, supplier data, or validated simulation package.",
        "Capture license and access rights before ingestion.",
        "Convert to machine-readable text with page, section, table, and figure preservation where possible.",
        "Extract candidate equations, variables, material data, propellant data, and assumptions.",
        "Run unit and dimensional checks on every candidate equation.",
        "Assign confidence, source status, and domain-owner review state.",
        "Store immutable source document, derived structured objects, embeddings, keyword indexes, and graph relationships.",
        "Promote only reviewed objects into solver-eligible status.",
    ])

    add_subsection(doc, "5.3 Knowledge Graph Relationships")
    add_monospace_block(
        doc,
        "Reference -> Document -> Section -> Equation -> Variables -> Units -> Dimensions\n"
        "Equation -> Assumptions -> Validity Range -> Verification Test -> Solver Function\n"
        "Material -> Manufacturing Process -> Allowables -> Structural Analysis Case\n"
        "Propellant -> Thermochemistry Model -> Combustion Solver -> Cooling Solver\n"
        "CAD Feature -> Solver Output -> Requirement -> Evidence Packet -> Human Review"
    )

    add_section(doc, "6. Engineering Computation Kernel")
    add_para(
        doc,
        "The computation kernel is the deterministic heart of COSMOS. It shall transform a design contract "
        "into a coherent design state by using versioned, tested, traceable physics and systems solvers. "
        "The AI/RAG layer can explain or retrieve; it cannot replace solver validation."
    )
    add_table(
        doc,
        ["Solver Family", "Responsibilities", "Early Maturity Target"],
        [
            ["Thermochemistry", "CEA/RocketCEA/Cantera integration, propellant properties, mixture ratio sweeps, combustion products", "Validated external-engine adapter"],
            ["Performance", "Thrust, Isp, c-star, expansion ratio, mass flow, pressure ratios", "Preliminary design capability"],
            ["Chamber and nozzle", "Chamber volume, contraction ratio, throat, bell contour, expansion, contour families", "Parametric geometry-ready outputs"],
            ["Injector", "Injector type selection, orifice sizing, pressure drop, spray assumptions, face layout", "Initial liquid/liquid injector generator"],
            ["Cooling", "Regenerative channel sizing, heat flux correlation, pressure drop, wall temperature estimates", "Conservative steady-state model"],
            ["Structure", "Hoop stress, thermal stress, buckling, margins, fatigue pre-checks", "Pre-FEA screening"],
            ["Ignition", "Igniter sizing basis, energy requirement, interface geometry", "Conceptual generator with review gate"],
            ["Feed dynamics", "Tank pressure, line losses, valve/transient model hooks", "Simscape/Modelica adapter initially"],
            ["Reliability", "FMEA, hazards, margin tracking, failure-mode catalog", "Design review support"],
            ["Optimization", "Parametric MDO, topology candidate loops, uncertainty sweeps", "OpenMDAO-first orchestration"],
        ],
        [1600, 5000, 2760],
        font_size=8.8,
    )

    add_subsection(doc, "6.1 Design State Object")
    add_para(
        doc,
        "A central EngineDesign object shall be the source of truth for a generated engine. It should be "
        "serializable, versioned, hashable, and replayable. It must include requirements, assumptions, solver "
        "inputs, solver outputs, CAD parameters, analysis packages, review status, and manufacturing-release "
        "evidence."
    )
    add_bullets(doc, [
        "Requirements: thrust, burn time, chamber pressure, propellant pair, mission point, safety factors.",
        "Design choices: cycle, mixture ratio, cooling method, injector family, material system, manufacturing process.",
        "Solver outputs: mass flow, throat area, expansion ratio, wall heat flux, channel dimensions, pressure drops, stresses.",
        "Geometry parameters: chamber length, radii, throat contour, nozzle control points, channel path, injector layout.",
        "Analysis records: mesh metadata, load cases, boundary conditions, convergence, result fields, comparison deltas.",
        "Release records: reviewer identity, open issues, waivers, signoff status, exported files, checksums.",
    ])

    add_section(doc, "7. CAD and Geometry Generation")
    add_para(
        doc,
        "COSMOS should generate parametric CAD first. Direct mesh generation is not enough for engineering "
        "design control. A parametric model preserves design intent, lets users export to Siemens NX, Ansys, "
        "Nastran, Fluent, OpenFOAM, and manufacturing systems, and enables future regeneration when requirements "
        "or solver models change."
    )
    add_table(
        doc,
        ["Geometry Object", "Generation Strategy", "Primary Export"],
        [
            ["Thrust chamber", "Parametric chamber barrel, contraction, throat, nozzle contour, wall thickness rules", "STEP, Parasolid if licensed, neutral mesh"],
            ["Nozzle", "Bell, conical, contour families; shape optimization around control points", "STEP, IGES, CFD surface mesh"],
            ["Regen channels", "Parametric helical/axial channels, manifolds, rib thickness, pressure-drop constraints", "STEP, analysis mesh, manufacturing report"],
            ["Injector head", "Pattern generator for element families, manifolds, faceplate, pressure-drop metadata", "STEP, hole table, manufacturing drawing package"],
            ["Igniter", "Parameterized igniter chamber, ports, interfaces, feed lines, mounting", "STEP, analysis-ready solid"],
            ["Manifolds", "Implicit/parametric flow distribution geometry, topology-assisted reinforcement", "STEP, CFD/FEA cases"],
            ["Mounts and interfaces", "Envelope-aware brackets, flanges, seals, fasteners, instrument ports", "STEP, Nastran/Abaqus structural model"],
        ],
        [1700, 5200, 2460],
        font_size=8.8,
    )

    add_subsection(doc, "7.1 Evidence-Bound CAD Feature Graph")
    add_para(
        doc,
        "Every CAD feature should carry metadata. For example, a throat radius feature should know the "
        "requirement it satisfies, the solver equation that produced it, the material and wall-temperature "
        "limit it must respect, the analysis cases that evaluated it, and the reviewer who accepted or rejected "
        "it. This feature-level traceability is one of the strongest unique COSMOS differentiators."
    )
    add_monospace_block(
        doc,
        "CADFeature(\n"
        "  feature_id='nozzle.throat.blend.radius',\n"
        "  generated_by='geometry.nozzle.v0_2',\n"
        "  source_requirements=['REQ-THRUST-001', 'REQ-PC-001'],\n"
        "  solver_outputs=['SOL-GASDYN-THROAT-AREA-004'],\n"
        "  evidence=['EQ-CHOKED-FLOW-001', 'MAT-INCONEL718-ALLOWABLE-003'],\n"
        "  analysis_cases=['FEA-THERMAL-021', 'CFD-NOZZLE-009'],\n"
        "  review_status='human_review_required'\n"
        ")"
    )

    add_section(doc, "8. Topology and Shape Optimization")
    add_para(
        doc,
        "Topology optimization shall be implemented as a candidate-generation and design-improvement layer, "
        "not as a direct manufacturing authority. It is most valuable for support structures, jackets, manifolds, "
        "mounts, brackets, injector structural supports, and additive-manufacturing lightweighting. Nozzle "
        "contours, cooling-channel paths, throat blends, and wall thickness should primarily use parametric or "
        "shape optimization because manufacturability and fluid/thermal constraints are tightly coupled."
    )
    add_table(
        doc,
        ["Optimization Type", "Best Use", "Release Rule"],
        [
            ["Topology optimization", "Mass reduction, stiffness, thermal paths, structural supports, manifolds, AM lattices", "Must be reconstructed into clean CAD and reanalyzed"],
            ["Shape optimization", "Nozzle contour, throat blend, injector passages, channel dimensions", "Must preserve constraints and regenerate parametric geometry"],
            ["Parametric optimization", "Mixture ratio, expansion ratio, chamber length, channel count, wall thickness", "Runs inside design-state/MDO workflow"],
            ["Lattice/implicit optimization", "AM thermal/structural tailoring, jacket reinforcement, support skins", "Requires process-specific manufacturability checks"],
            ["Robust optimization", "Uncertainty margins for material, thermal, pressure, manufacturing tolerance", "Must report sensitivity and confidence intervals"],
        ],
        [1850, 4750, 2760],
        font_size=9,
    )
    add_subsection(doc, "8.1 Topology Loop")
    add_numbers(doc, [
        "Generate conservative parametric baseline geometry.",
        "Define preserve regions, loads, boundary conditions, manufacturing constraints, thermal paths, and no-go volumes.",
        "Run topology or shape optimization with explicit objective and constraints.",
        "Convert raw output into clean parametric or implicit CAD.",
        "Run geometry checks, wall-thickness checks, AM overhang/support checks, and mesh-quality checks.",
        "Re-run external CFD/FEA/thermal analysis.",
        "Compare optimized result against baseline in COSMOS.",
        "Require human review before release.",
    ])

    add_section(doc, "9. External Simulation and Dynamic Validation")
    add_para(
        doc,
        "COSMOS should be solver-neutral and export complete analysis packages. External software results "
        "should be uploaded back into COSMOS for side-by-side comparison against internal predictions. This "
        "secondary review loop is essential for credibility before any 3D-printing release."
    )
    add_table(
        doc,
        ["Validation Domain", "Recommended Tools", "COSMOS Integration"],
        [
            ["CAD and manufacturing review", "Siemens NX, FreeCAD, nTop, CAM systems", "STEP/Parasolid import/export, feature metadata, manufacturing notes"],
            ["CFD and combustion flow", "Ansys Fluent, OpenFOAM, STAR-CCM+, SU2", "Mesh, boundary conditions, residuals, fields, heat flux, pressure drop"],
            ["Structural FEA/FEM", "Ansys Mechanical/Workbench, MSC Nastran, Abaqus", "Loads, constraints, thermal maps, stress, strain, safety factor, modal data"],
            ["Explicit/transient dynamics", "LS-DYNA, Abaqus Explicit", "Shock, burst, impact, severe transient load cases"],
            ["Multibody dynamics", "MSC Adams, Simcenter 3D Motion", "Gimbal, actuator, valve, mechanism, mounting loads"],
            ["System dynamics", "MATLAB Simulink/Simscape, Dymola/Modelica", "Feed-system transients, pressurization, valve dynamics, control loops"],
            ["Optimization/UQ", "OpenMDAO, Dakota, optiSLang, modeFRONTIER", "MDO, sensitivity, uncertainty, design-space exploration"],
        ],
        [1750, 3300, 4310],
        font_size=8.8,
    )

    add_subsection(doc, "9.1 Imported Result Package Requirements")
    add_bullets(doc, [
        "External solver name, version, license environment, and date.",
        "Geometry file checksum and design-state version.",
        "Mesh count, mesh-quality metrics, units, coordinate system, and named selections.",
        "Boundary conditions, material model, turbulence/combustion/thermal model, load case, and solver settings.",
        "Residuals, convergence criteria, failed warnings, and reviewer notes.",
        "Primary outputs: temperatures, pressures, stresses, deformation, modal frequencies, heat flux, pressure drop, margin.",
        "Difference table comparing external results with COSMOS predictions.",
    ])

    add_subsection(doc, "9.2 Dynamic Simulation Recommendation")
    add_para(
        doc,
        "For dynamic simulation, COSMOS should start with system-level transients and mechanism dynamics before "
        "attempting high-fidelity transient combustion. Simulink/Simscape and Modelica/Dymola are strong for "
        "feed-system dynamics, pressurization, and controls. Adams or Simcenter Motion are strong for gimbals, "
        "valves, actuators, and moving assemblies. LS-DYNA or Abaqus Explicit should be reserved for severe "
        "transient structural events and nonlinear dynamics."
    )

    add_section(doc, "10. UI/UX Architecture")
    add_para(
        doc,
        "The UI should feel like a professional engineering workbench rather than a chatbot or calculator. "
        "The user should see requirements, evidence, geometry, solver status, external validation, and release "
        "readiness as one coherent digital thread."
    )
    add_table(
        doc,
        ["Workspace", "Purpose", "Key UI Elements"],
        [
            ["Dashboard", "Project overview and release readiness", "Design status, open risks, latest run, review passport"],
            ["Design Contract", "Requirements and constraints", "Validated forms, unit-aware inputs, requirement IDs"],
            ["Evidence Browser", "RAG and knowledge foundation", "Citations, equations, assumptions, validity ranges"],
            ["Sizing and Solvers", "Engineering computation", "Run graph, solver logs, convergence, warnings, margins"],
            ["Geometry Studio", "CAD generation and topology candidates", "3D viewer, feature tree, parametric controls, export panel"],
            ["Simulation Hub", "External case management", "Tool adapters, case packages, upload results, mesh metadata"],
            ["Comparison Cockpit", "Side-by-side validation", "COSMOS vs external result tables, plots, deltas, pass/fail"],
            ["Documentation Center", "Engineering reports and design basis", "Templates, generated docs, versioned outputs"],
            ["Release Gate", "Human review and manufacturing release", "Checklist, issue log, signoff, export controls"],
        ],
        [1600, 3500, 4260],
        font_size=8.8,
    )
    add_subsection(doc, "10.1 UI Principles")
    add_bullets(doc, [
        "Desktop-first, local-first, professional engineering layout.",
        "Dense but readable workspaces; avoid marketing-style screens in operational views.",
        "All numerical inputs are unit-aware and validation-aware.",
        "Every output has an evidence or computation trail accessible from the UI.",
        "The AI assistant appears as an engineering copilot and review explainer, not as the hidden authority.",
        "3D geometry, plots, tables, solver logs, and citations should be visible together when relevant.",
        "Release readiness must be visible as a formal state, not implied by successful geometry generation.",
    ])

    add_section(doc, "11. Engineering Documentation System")
    add_para(
        doc,
        "COSMOS shall generate engineering documentation automatically and store it as traceable design evidence. "
        "Documentation is not a side output; it is part of the design authority. Every report should be generated "
        "from the design state, source evidence, solver outputs, CAD metadata, external analysis results, and review "
        "records."
    )
    add_table(
        doc,
        ["Document Type", "Contents", "Release Role"],
        [
            ["Design Basis Document", "Requirements, assumptions, selected architecture, propellants, materials, validity ranges", "Defines why the engine exists and what it must satisfy"],
            ["Calculation Package", "Equations, variables, constants, solver versions, numerical outputs, margins", "Auditable engineering computation record"],
            ["CAD Generation Report", "Feature tree, parametric values, generator version, geometry checks, export files", "Shows how geometry was produced"],
            ["Topology Optimization Report", "Objectives, constraints, preserve regions, candidate comparison, reconstruction notes", "Prevents raw topology output from becoming unreviewed design truth"],
            ["CFD Report", "Mesh, BCs, solver model, residuals, fields, comparison deltas", "Fluid/thermal validation evidence"],
            ["FEA/FEM Report", "Loads, constraints, material model, stress/strain/thermal results, margins", "Structural validation evidence"],
            ["Manufacturing Release Passport", "All required gates, signoffs, file checksums, open waivers, print/process notes", "Final release-control artifact"],
            ["Change Impact Report", "Before/after requirements, geometry, solver outputs, analysis deltas", "Controls revisions and regression risk"],
        ],
        [1800, 5000, 2560],
        font_size=8.7,
    )
    add_subsection(doc, "11.1 Documentation Storage and Traceability")
    add_bullets(doc, [
        "Every generated document receives a document ID, version, content hash, source design-state hash, and authoring pipeline version.",
        "Documents remain immutable after release; revisions create new versions.",
        "Reports shall be exportable as DOCX, PDF, Markdown, JSON evidence bundles, and project archive records.",
        "Every report section should be traceable to source data or generated computation.",
        "External analysis reports uploaded by reviewers should be stored beside generated COSMOS reports and compared in the UI.",
    ])

    add_section(doc, "12. Data, Persistence, and Digital Thread")
    add_para(
        doc,
        "COSMOS needs a digital-thread storage model that can operate locally but later scale to team and "
        "enterprise use. The first implementation can use SQLite/DuckDB plus file-based artifact storage; "
        "the architecture should leave a path to PostgreSQL, object storage, and vector/graph services."
    )
    add_table(
        doc,
        ["Data Store", "Recommended Initial Choice", "Future Enterprise Option"],
        [
            ["Project metadata", "SQLite", "PostgreSQL"],
            ["Analytical tabular results", "DuckDB/SQLite", "PostgreSQL/DuckDB server"],
            ["Large arrays and fields", "HDF5/Zarr", "Object storage plus metadata DB"],
            ["CAD artifacts", "Versioned files with checksums", "PDM/PLM integration"],
            ["RAG vectors", "FAISS/Qdrant local", "Qdrant/pgvector managed service"],
            ["Knowledge graph", "NetworkX/SQLite initially", "Neo4j/PostgreSQL graph extensions"],
            ["Generated documents", "File store plus document table", "Document management integration"],
            ["Audit logs", "Append-only JSONL/SQLite", "Immutable audit service"],
        ],
        [2200, 3600, 3560],
        font_size=9,
    )

    add_section(doc, "13. Security, IP, and Compliance Controls")
    add_para(
        doc,
        "COSMOS should be treated as company intellectual property from the beginning. Source code, schemas, "
        "solver logic, CAD generators, topology workflows, curated datasets, generated reports, UI workflows, "
        "and validation pipelines are proprietary company assets unless explicitly licensed otherwise."
    )
    add_table(
        doc,
        ["Control Area", "Requirement"],
        [
            ["Ownership", "All contributors must have clear invention assignment, contribution, and confidentiality terms."],
            ["Third-party code", "Track licenses and never mix incompatible code into proprietary modules."],
            ["Reference data", "Store license, source, access rights, and allowed-use status for every ingested source."],
            ["Trade secrets", "Protect solver workflows, architecture pattern, model calibration data, and CAD generator logic."],
            ["Access control", "Role-based access for source data, projects, release packages, and export functions."],
            ["Auditability", "Record who generated, changed, reviewed, exported, or released every design artifact."],
            ["Export control", "Flag propulsion design data for jurisdiction-specific legal review before external sharing."],
            ["Clean-room posture", "Do not reverse-engineer external proprietary computational-engineering products."],
        ],
        [2100, 7260],
        font_size=9,
    )

    add_section(doc, "14. API and Interoperability Strategy")
    add_para(
        doc,
        "COSMOS shall interoperate with major engineering tools through neutral formats and adapter packages. "
        "The goal is not to replace every high-fidelity solver immediately, but to generate complete, traceable "
        "case packages and to ingest their results."
    )
    add_table(
        doc,
        ["Artifact", "Preferred Formats"],
        [
            ["CAD solids", "STEP AP242, STEP AP214, IGES, Parasolid where licensed"],
            ["Mesh", "CGNS, Gmsh MSH, SU2, VTK/VTU, Nastran BDF, Abaqus INP, Ansys CDB"],
            ["OpenFOAM", "OpenFOAM case folders with boundary, system, constant directories"],
            ["CFD/FEA results", "VTK/VTU, CGNS, CSV summaries, HDF5/Zarr field data, JSON metadata"],
            ["Print preparation", "3MF preferred, STL only as final mesh export after release gate"],
            ["Documentation", "DOCX, PDF, Markdown, JSON evidence bundle"],
            ["Project archive", "COSMOS project package with manifest, checksums, versions, and dependency lock"],
        ],
        [2300, 7060],
        font_size=9,
    )

    add_section(doc, "15. Quality, Verification, and Scientific Credibility")
    add_para(
        doc,
        "The software cannot become world-class unless its validation is treated as a first-class subsystem. "
        "COSMOS must distinguish between conceptual design, preliminary design, analysis-ready geometry, "
        "externally validated design, and manufacturing release. Passing a calculation is not the same as "
        "being safe to manufacture."
    )
    add_table(
        doc,
        ["Quality Gate", "Required Evidence"],
        [
            ["Code quality", "Compile, lint, type-check, unit tests, integration tests, coverage trend"],
            ["Numerical safety", "Finite-number checks, units, dimensions, bounds, operating ranges"],
            ["Equation verification", "Known examples, source reproduction, independent derivation checks"],
            ["Solver validation", "Benchmark data, external solver comparison, uncertainty ranges"],
            ["CAD validation", "Watertight solids, feature integrity, wall thickness, interference, manufacturability"],
            ["Mesh validation", "Mesh quality, named boundaries, unit consistency, reproducible case generation"],
            ["External review", "Imported results, comparison deltas, reviewer comments, accepted waivers"],
            ["Release control", "Open issue closure, human signoff, artifact checksums, final export record"],
        ],
        [1900, 7460],
        font_size=9,
    )

    add_subsection(doc, "15.1 Release State Model")
    add_table(
        doc,
        ["State", "Meaning"],
        [
            ["DRAFT", "Requirements or knowledge are incomplete."],
            ["CONCEPTUAL", "Preliminary sizing exists, but geometry or analysis is incomplete."],
            ["ANALYSIS_READY", "Geometry and case packages are ready for external analysis."],
            ["ANALYSIS_RUNNING", "External or internal simulations are active."],
            ["ANALYSIS_FAILED", "One or more validation cases failed or did not converge."],
            ["REVIEW_REQUIRED", "Automated checks passed enough for human review."],
            ["RELEASE_BLOCKED", "Open risks, missing evidence, or failed checks prevent manufacturing export."],
            ["MANUFACTURING_RELEASED", "Human signoff completed and release passport generated."],
        ],
        [2200, 7160],
        font_size=9,
    )

    add_section(doc, "16. Development Roadmap")
    add_para(
        doc,
        "The roadmap should avoid trying to build the full vision at once. A credible path starts with one "
        "narrow engine family, one propellant pair, one parametric geometry generator, one external simulation "
        "round trip, and one documentation package."
    )
    add_table(
        doc,
        ["Phase", "Scope", "Exit Criteria"],
        [
            ["0 - Foundation hardening", "Fix package imports, compile errors, validation gaps, project tooling, canonical data", "Full test suite, compileall, ruff, mypy baseline pass"],
            ["1 - Knowledge/RAG MVP", "Reference/document/equation ingestion, vector+keyword search, citation answers", "Approved equations can be retrieved with source and validity data"],
            ["2 - Preliminary engine sizing", "Propellant pair, thrust, Pc, burn time, chamber/nozzle sizing", "Deterministic design state and calculation report"],
            ["3 - Parametric CAD MVP", "Chamber, throat, nozzle, simple injector placeholder, STEP export", "Regeneratable CAD with feature metadata"],
            ["4 - External analysis round trip", "OpenFOAM/Fluent and Nastran/Ansys package export/import", "Comparison cockpit with uploaded results"],
            ["5 - Regen cooling and injector depth", "Channel geometry, pressure drop, heat flux, injector face layout", "Analysis-ready chamber/injector assemblies"],
            ["6 - Optimization layer", "OpenMDAO parametric optimization and topology candidate workflow", "Baseline vs optimized comparison and reconstruction report"],
            ["7 - Release passport", "Human review workflow, documentation center, manufacturing export controls", "No print export without completed release gate"],
            ["8 - Enterprise maturity", "Team database, access control, plugin SDK, validation dataset management", "Multi-user digital thread and audit capability"],
        ],
        [1350, 4750, 3260],
        font_size=8.7,
    )

    add_section(doc, "17. Immediate Repo-Level Actions")
    add_bullets(doc, [
        "Fix knowledge package path mismatch between repository tests and implementation.",
        "Repair or quarantine the non-compiling Quantity model before it enters the active RAG schema.",
        "Promote one propellant database as canonical and ensure the default loader points to it.",
        "Add pyproject.toml, dependency lock, formatter/linter config, and CI commands.",
        "Implement shared finite numeric validation and apply it across propellants, variables, units, constants, and materials.",
        "Remove debug print statements from models and prevent import-time side effects in configuration/settings.",
        "Create first design-state schema before writing more CAD or solver features.",
    ])

    add_section(doc, "18. Appendix A - Initial COSMOS Module Map")
    add_monospace_block(
        doc,
        "cosmos/\n"
        "  core/                  units, constants, validation, config, logging\n"
        "  knowledge/             RAG schemas, ingestion, graph, retrieval, approval\n"
        "  physics/               equations and scientific models only\n"
        "  systems/               chamber, nozzle, injector, cooling, igniter, feed\n"
        "                         structure and reliability assemblies\n"
        "  backend/               orchestration, workflows, APIs, project state\n"
        "  geometry/              CAD generators, feature graph, export adapters\n"
        "  optimization/          OpenMDAO, topology/shape loops, UQ, sensitivity\n"
        "  simulation_adapters/   Fluent, OpenFOAM, Nastran, Ansys, Abaqus, NX, Simulink\n"
        "  validation/            benchmarks, external comparison, release gates\n"
        "  documentation/         report templates, design basis, release passport\n"
        "  gui/                   Qt workbench, 3D viewer, comparison cockpit\n"
        "  databases/             project, knowledge, material, propellant, result stores\n"
        "  tests/                 unit, integration, regression, validation datasets"
    )

    add_section(doc, "19. Appendix B - Human Review Checklist")
    add_bullets(doc, [
        "Design contract complete and approved.",
        "All equations and material/property data are source-bound and in approved status.",
        "All solver inputs are finite, unit-valid, dimension-valid, and within validity range.",
        "CAD solids pass geometry integrity and manufacturability checks.",
        "External CFD result package uploaded and compared.",
        "External FEA/FEM result package uploaded and compared.",
        "Thermal and cooling margin reviewed.",
        "Structural margin, modal/vibration risk, and fatigue assumptions reviewed.",
        "Topology reconstruction, if used, was reanalyzed after CAD cleanup.",
        "Open risks, assumptions, and waivers are explicitly accepted or closed.",
        "Manufacturing process, material, inspection, and test-plan notes are attached.",
        "Human reviewer signs release passport before 3MF/STL manufacturing export.",
    ])

    add_section(doc, "20. Appendix C - Design Input/Output Contract")
    add_table(
        doc,
        ["Category", "Minimum Fields"],
        [
            ["User inputs", "Propellant pair, thrust, chamber pressure, burn time, environment, feed architecture, cooling method, material/process, safety factors"],
            ["Computed outputs", "Mass flow, mixture ratio, throat area, chamber dimensions, expansion ratio, nozzle contour, cooling channels, injector layout, wall thickness"],
            ["CAD outputs", "Chamber/nozzle assembly, regen jacket/channels, injector head, igniter, manifolds, mounts, interfaces"],
            ["CAE outputs", "CFD case, FEA/FEM case, thermal case, dynamics/system model package, mesh files, result metadata"],
            ["Documentation outputs", "Design basis, calculation package, CAD report, analysis comparison, review passport, change report"],
            ["Manufacturing outputs", "STEP/Parasolid, drawings/tables, 3MF/STL after release, inspection notes, process assumptions"],
        ],
        [2200, 7160],
        font_size=9,
    )

    add_section(doc, "21. Closing Direction")
    add_para(
        doc,
        "COSMOS should be built as a proprietary evidence-driven computational engineering system. Its unique "
        "value should come from the integration of traceable RAG knowledge, deterministic engineering solvers, "
        "feature-level CAD provenance, topology-aware design improvement, external solver comparison, generated "
        "engineering documentation, and controlled human release. That combination is the company-owned pattern. "
        "The immediate task is to harden the foundation, then build one narrow end-to-end engine workflow that "
        "proves the TRACE-GEN loop from requirements to reviewed CAD and documentation."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
